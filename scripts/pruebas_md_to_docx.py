"""
Genera PRUEBAS.docx desde PRUEBAS.md (tablas Word nativas).

Uso (desde la raíz del repo):
  pip install python-docx
  python scripts/pruebas_md_to_docx.py

Salida por defecto: PRUEBAS.docx junto a PRUEBAS.md
"""

from __future__ import annotations

import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.shared import Pt
except ImportError:
    print("Instala: pip install python-docx", file=sys.stderr)
    sys.exit(1)


def repo_root() -> Path:
    return Path(__file__).resolve().parents[1]


def split_table_row(line: str) -> list[str]:
    parts = [p.strip() for p in line.strip().split("|")]
    if parts and parts[0] == "":
        parts = parts[1:]
    if parts and parts[-1] == "":
        parts = parts[:-1]
    return parts


def is_table_separator(cells: list[str]) -> bool:
    if not cells:
        return False
    return all(re.fullmatch(r"[-:\s]+", c or "") for c in cells)


def strip_md_inline(s: str) -> str:
    s = re.sub(r"\*\*(.+?)\*\*", r"\1", s)
    s = re.sub(r"`([^`]+)`", r"\1", s)
    s = re.sub(r"\*(.+?)\*", r"\1", s)
    return s


def add_paragraph_plain(doc: Document, text: str) -> None:
    text = strip_md_inline(text.strip())
    if not text:
        return
    p = doc.add_paragraph()
    p.add_run(text)
    p.paragraph_format.space_after = Pt(6)


def add_table(doc: Document, rows: list[list[str]], header: bool = True) -> None:
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j in range(ncols):
            cell_text = strip_md_inline(row[j]) if j < len(row) else ""
            cell = table.rows[i].cells[j]
            cell.text = cell_text
            for para in cell.paragraphs:
                para.paragraph_format.space_after = Pt(0)
            if header and i == 0:
                for run in cell.paragraphs[0].runs:
                    run.bold = True


def parse_md_to_docx(md_path: Path, out_path: Path) -> None:
    lines = md_path.read_text(encoding="utf-8").splitlines()
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    i = 0
    while i < len(lines):
        line = lines[i]

        if line.strip() == "---":
            i += 1
            continue

        if line.startswith("# ") and not line.startswith("## "):
            doc.add_heading(line[2:].strip(), level=0)
            i += 1
            continue

        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
            i += 1
            continue

        if line.startswith("|"):
            block: list[str] = []
            while i < len(lines) and lines[i].startswith("|"):
                block.append(lines[i])
                i += 1
            rows_raw = [split_table_row(b) for b in block]
            rows: list[list[str]] = []
            for r in rows_raw:
                if is_table_separator(r):
                    continue
                rows.append(r)
            if rows:
                add_table(doc, rows, header=True)
            doc.add_paragraph()
            continue

        if line.strip():
            if line.strip().startswith("*") and line.strip().endswith("*") and line.count("*") >= 2:
                p = doc.add_paragraph()
                run = p.add_run(strip_md_inline(line.strip().strip("*")))
                run.italic = True
                p.paragraph_format.space_after = Pt(6)
            else:
                add_paragraph_plain(doc, line)
        i += 1

    doc.save(out_path)
    print(f"Escrito: {out_path}")


def main() -> None:
    root = repo_root()
    md = root / "PRUEBAS.md"
    out = root / "PRUEBAS.docx"
    if not md.is_file():
        print(f"No existe {md}", file=sys.stderr)
        sys.exit(1)
    parse_md_to_docx(md, out)


if __name__ == "__main__":
    main()

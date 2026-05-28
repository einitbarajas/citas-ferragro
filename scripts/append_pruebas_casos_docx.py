"""
Añade al PRUEBAS.docx la sección 8: un caso de prueba tabulado por cada test pytest
(formato Field | Description, encabezado verde).

Uso (desde la raíz del repo):
  python scripts/pruebas_md_to_docx.py docs/PRUEBAS.md docs/PRUEBAS.docx
  python scripts/append_pruebas_casos_docx.py
"""

from __future__ import annotations

import ast
import subprocess
import sys
from dataclasses import dataclass
from pathlib import Path

try:
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt
except ImportError:
    print("Instala: pip install python-docx", file=sys.stderr)
    sys.exit(1)


ROOT = Path(__file__).resolve().parents[1]
BACKEND = ROOT / "backend"
DOCX_PATH = ROOT / "docs" / "PRUEBAS.docx"
HEADER_FILL = "C6E0B4"  # verde claro (estilo plantilla)


@dataclass
class PytestCase:
    number: int
    nodeid: str
    module_file: str
    func_name: str
    param: str | None
    docstring: str | None


FILE_META: dict[str, dict[str, str]] = {
    "test_api_system_smoke.py": {
        "priority": "Alta",
        "layer": "API HTTP",
        "pre": "Backend importable; no requiere PostgreSQL.",
    },
    "test_db_crud_functions.py": {
        "priority": "Alta",
        "layer": "PostgreSQL / PL-pgSQL",
        "pre": "BD `db_trabajo` con esquema y CRUD (`db\\run-database-all.ps1`); `DATABASE_URL` en `.env`.",
    },
    "test_appointment_service_db.py": {
        "priority": "Alta",
        "layer": "Backend + PostgreSQL",
        "pre": "BD preparada; bodega/proveedor de prueba (seed o datos creados por tests CRUD).",
    },
    "test_franjas_per_team_logic.py": {
        "priority": "Alta",
        "layer": "Backend + PostgreSQL",
        "pre": "BD preparada; equipos de descarga configurados.",
    },
    "test_unload_team_resolve.py": {
        "priority": "Media",
        "layer": "Backend + PostgreSQL",
        "pre": "BD preparada.",
    },
    "test_unload_team_names.py": {
        "priority": "Media",
        "layer": "Backend + PostgreSQL",
        "pre": "BD preparada.",
    },
    "test_analytics_summary_timezone.py": {
        "priority": "Media",
        "layer": "Backend + PostgreSQL",
        "pre": "BD preparada; zona `America/Bogota`.",
    },
    "test_logistics_business_rules.py": {
        "priority": "Alta",
        "layer": "Backend (unitario)",
        "pre": "Sin PostgreSQL; mocks de sesión DB.",
    },
    "test_appointment_service_unit.py": {
        "priority": "Alta",
        "layer": "Backend (unitario)",
        "pre": "Sin PostgreSQL.",
    },
    "test_appointment_windows_unit.py": {
        "priority": "Media",
        "layer": "Backend (unitario)",
        "pre": "Sin PostgreSQL.",
    },
    "test_notification_service_unit.py": {
        "priority": "Media",
        "layer": "Backend (unitario)",
        "pre": "Sin PostgreSQL.",
    },
    "test_range_bounds_unit.py": {
        "priority": "Media",
        "layer": "Backend (unitario)",
        "pre": "Sin PostgreSQL.",
    },
    "test_email_utils.py": {
        "priority": "Media",
        "layer": "Backend / correo",
        "pre": "Sin PostgreSQL.",
    },
    "test_smtp_profile_config.py": {
        "priority": "Media",
        "layer": "Configuración SMTP",
        "pre": "Sin PostgreSQL; variables de entorno de prueba en el test.",
    },
}

DEFAULT_META = {
    "priority": "Media",
    "layer": "Backend",
    "pre": "Ver §3 (pytest con o sin PostgreSQL según archivo).",
}


def set_cell_shading(cell, fill_hex: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def humanize_test_name(func_name: str) -> str:
    name = func_name.removeprefix("test_").replace("_", " ")
    return name[:1].upper() + name[1:] if name else func_name


def collect_pytest_cases() -> list[PytestCase]:
    py = BACKEND / ".venv" / "Scripts" / "python.exe"
    if not py.is_file():
        py = Path(sys.executable)
    env = {**dict(**__import__("os").environ), "PYTHONPATH": str(BACKEND)}
    proc = subprocess.run(
        [str(py), "-m", "pytest", "tests/", "--collect-only", "-q"],
        cwd=BACKEND,
        capture_output=True,
        text=True,
        env=env,
    )
    if proc.returncode != 0:
        print(proc.stderr or proc.stdout, file=sys.stderr)
        sys.exit(proc.returncode)

    lines = [ln.strip() for ln in proc.stdout.splitlines() if ln.strip().startswith("tests/")]
    cases: list[PytestCase] = []
    for i, line in enumerate(lines, start=1):
        nodeid = line.split()[0] if " " in line else line
        path_part, _, rest = nodeid.partition("::")
        module_file = Path(path_part).name
        if "[" in rest:
            func_name, _, param_part = rest.partition("[")
            param = param_part.rstrip("]")
        else:
            func_name = rest
            param = None
        cases.append(
            PytestCase(
                number=i,
                nodeid=nodeid,
                module_file=module_file,
                func_name=func_name,
                param=param,
                docstring=_load_test_docstring(module_file, func_name),
            )
        )
    return cases


def _load_test_docstring(module_file: str, func_name: str) -> str | None:
    path = BACKEND / "tests" / module_file
    if not path.is_file():
        return None
    try:
        tree = ast.parse(path.read_text(encoding="utf-8"))
    except SyntaxError:
        return None
    for node in tree.body:
        if isinstance(node, ast.FunctionDef) and node.name == func_name:
            return ast.get_docstring(node)
    return None


def _infer_expected(case: PytestCase) -> str:
    if case.docstring:
        first = case.docstring.strip().splitlines()[0]
        return f"{first} (aserciones pytest deben cumplirse)."
    return "Todas las aserciones del test pasan sin excepción (pytest PASSED)."


def _test_title(case: PytestCase) -> str:
    title = humanize_test_name(case.func_name)
    if case.param:
        title = f"{title} — {case.param}"
    return title


def _test_data(case: PytestCase) -> str:
    meta = FILE_META.get(case.module_file, DEFAULT_META)
    if "PostgreSQL" in meta["pre"] or "BD" in meta["pre"]:
        return "DATABASE_URL apuntando a `db_trabajo`; datos según seed o fixtures del test."
    if case.param:
        return f"Parámetro: {case.param}"
    return "Datos mock / constantes definidos en el código del test (sin BD)."


def add_field_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Field"
    hdr[1].text = "Description"
    for cell in hdr:
        set_cell_shading(cell, HEADER_FILL)
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for i, (field, desc) in enumerate(rows, start=1):
        table.rows[i].cells[0].text = field
        table.rows[i].cells[1].text = desc
    doc.add_paragraph()


def append_test_cases(doc: Document, cases: list[PytestCase]) -> None:
    doc.add_page_break()
    doc.add_heading("8. Casos de prueba detallados (pytest)", level=1)
    intro = doc.add_paragraph()
    intro.add_run(
        "Cada fila siguiente documenta un test automatizado del inventario §3. "
        "Formato alineado a plantilla de casos de prueba (Field / Description). "
        f"Total: {len(cases)} casos."
    )
    intro.paragraph_format.space_after = Pt(12)

    current_file = ""
    file_index = 0
    for case in cases:
        if case.module_file != current_file:
            current_file = case.module_file
            file_index += 1
            doc.add_heading(f"8.{file_index} {current_file}", level=2)

        meta = FILE_META.get(case.module_file, DEFAULT_META)
        title = _test_title(case)
        doc.add_heading(f"Test case #{case.number}: {title}", level=3)

        cmd = f'cd backend\n$env:PYTHONPATH="."\npy -m pytest {case.nodeid} -v'

        rows = [
            ("Test Case #", str(case.number)),
            ("Test Priority", meta["priority"]),
            ("Test Title/Name", title),
            ("Test Summary", case.docstring or f"Validación automatizada: {humanize_test_name(case.func_name)}."),
            ("Test Steps", f"1. Preparar entorno ({meta['pre']})\n2. Ejecutar:\n{cmd}\n3. Verificar salida: 1 passed."),
            ("Test Data", _test_data(case)),
            ("Expected Result", _infer_expected(case)),
            (
                "Actual Result",
                "Automatizado: resultado de la última ejecución pytest (local o CI). Consultar §3 para suite completa.",
            ),
            ("Status", "Automatizado (pytest)"),
            ("Pre-condition", meta["pre"]),
            (
                "Post-condition",
                "Sin efectos persistentes en BD (tests con PostgreSQL usan rollback) o sin cambios de estado (unitarios).",
            ),
            (
                "Notes/Comments",
                f"Archivo: {case.module_file} | Capa: {meta['layer']} | Node ID: {case.nodeid}",
            ),
        ]
        add_field_table(doc, rows)


def rebuild_base_docx() -> None:
    py = BACKEND / ".venv" / "Scripts" / "python.exe"
    if not py.is_file():
        py = Path(sys.executable)
    script = ROOT / "scripts" / "pruebas_md_to_docx.py"
    subprocess.run([str(py), str(script), "docs/PRUEBAS.md", "docs/PRUEBAS.docx"], cwd=ROOT, check=True)


def main() -> None:
    rebuild_base_docx()
    cases = collect_pytest_cases()
    doc = Document(str(DOCX_PATH))
    append_test_cases(doc, cases)
    doc.save(DOCX_PATH)
    print(f"PRUEBAS.docx generado: base + sección 8 ({len(cases)} casos tabulados)")


if __name__ == "__main__":
    main()

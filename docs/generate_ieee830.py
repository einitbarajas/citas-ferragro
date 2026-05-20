# -*- coding: utf-8 -*-
"""
Genera ESPECIFICACION_REQUISITOS_IEEE830_FERRAGRO.docx (IEEE 830) desde cero.

Uso:
  pip install python-docx
  python docs/generate_ieee830.py
"""
from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

ROOT = Path(__file__).resolve().parent
OUTPUT = ROOT / "ESPECIFICACION_REQUISITOS_IEEE830_FERRAGRO.docx"
VERSION = "1.0"
TODAY = date.today().strftime("%d/%m/%Y")


class DocBuilder:
    def __init__(self) -> None:
        self.doc = Document()
        self._setup_styles()
        self._rf = 0
        self._rnf = 0

    def _setup_styles(self) -> None:
        section = self.doc.sections[0]
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2.5)
        normal = self.doc.styles["Normal"]
        normal.font.name = "Calibri"
        normal.font.size = Pt(11)

    def _p(self, text: str = "", *, bold: bool = False, size: int | None = None) -> None:
        p = self.doc.add_paragraph()
        run = p.add_run(text)
        if bold:
            run.bold = True
        if size:
            run.font.size = Pt(size)

    def h(self, text: str, level: int) -> None:
        self.doc.add_heading(text, level=level)

    def bullet(self, text: str) -> None:
        self.doc.add_paragraph(text, style="List Bullet")

    def numbered(self, text: str) -> None:
        self.doc.add_paragraph(text, style="List Number")

    def req_functional(
        self,
        title: str,
        description: str,
        *,
        priority: str = "Alta",
        actors: str = "Sistema",
        processing: str = "Validación en API, persistencia en PostgreSQL, respuesta JSON estándar.",
    ) -> None:
        self._rf += 1
        code = f"RF-{self._rf:03d}"
        self.h(f"{code}: {title}", 3)
        self.bullet(f"Prioridad: {priority}")
        self.bullet(f"Actores: {actors}")
        self.bullet(f"Descripción: {description}")
        self.bullet(f"Procesamiento: {processing}")

    def req_nonfunctional(
        self,
        title: str,
        description: str,
        *,
        category: str = "Calidad",
        metric: str | None = None,
    ) -> None:
        self._rnf += 1
        code = f"RNF-{self._rnf:03d}"
        self.h(f"{code}: {title}", 3)
        self.bullet(f"Categoría: {category}")
        self.bullet(f"Descripción: {description}")
        if metric:
            self.bullet(f"Métrica / criterio: {metric}")

    def cover(self) -> None:
        title = self.doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = title.add_run("ESPECIFICACIÓN DE REQUISITOS DE SOFTWARE")
        r.bold = True
        r.font.size = Pt(16)

        sub = self.doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = sub.add_run("Sistema Ferragro — Gestión de Citas de Entrega de Materiales")
        r2.bold = True
        r2.font.size = Pt(14)

        self.doc.add_paragraph()
        for line in (
            f"Versión del documento: {VERSION}",
            f"Fecha: {TODAY}",
            "Organización: Ferragro",
            "Estándar de referencia: IEEE Std 830-1998",
            "Estado: Borrador para revisión",
        ):
            p = self.doc.add_paragraph(line)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        self.doc.add_page_break()

    def build(self) -> Document:
        self.cover()
        self._section_1_intro()
        self._section_2_general()
        self._section_3_specific()
        self._section_4_traceability()
        return self.doc

    def _section_1_intro(self) -> None:
        self.h("1. Introducción", 1)
        self.h("1.1 Propósito", 2)
        self._p(
            "Este documento define los requisitos funcionales y no funcionales del software "
            "Ferragro: portal web para agendar, revisar y administrar citas de entrega de "
            "materiales entre proveedores y el personal interno de Ferragro (Admin y Logística). "
            "Sirve como contrato entre las partes interesadas y el equipo de desarrollo, y como "
            "base para pruebas de aceptación y mantenimiento."
        )

        self.h("1.2 Alcance", 2)
        self._p("El producto comprende:")
        self.bullet("Aplicación web (interfaz de usuario) accesible desde navegador.")
        self.bullet("API REST (backend) con autenticación JWT y sesiones de actualización.")
        self.bullet("Base de datos relacional PostgreSQL para credenciales, usuarios, proveedores, citas, franjas horarias y auditoría.")
        self._p("Queda fuera de alcance, salvo integración futura explícita:")
        self.bullet("Sistemas ERP o inventario de Ferragro.")
        self.bullet("Facturación, pagos en línea y firma electrónica de documentos legales.")
        self.bullet("Aplicaciones móviles nativas (iOS/Android); el acceso móvil se realiza vía navegador responsivo.")

        self.h("1.3 Definiciones, acrónimos y abreviaturas", 2)
        terms = [
            ("API", "Interfaz de programación de aplicaciones; servicio HTTP JSON del backend."),
            ("Admin", "Usuario interno con permisos completos de configuración y eliminación."),
            ("Cita", "Reserva de una ventana horaria para entrega de material en instalaciones Ferragro."),
            ("Credencial", "Registro de correo y contraseña (hash) usado para autenticación."),
            ("Franja horaria", "Intervalo de tiempo en el que se permiten citas (configurable por día)."),
            ("JWT", "JSON Web Token; credencial de acceso de corta duración."),
            ("Logística", "Usuario interno operativo: revisión de citas, proveedores y reportes."),
            ("NIT", "Número de identificación tributaria del proveedor (10 dígitos en el sistema)."),
            ("Proveedor", "Empresa externa que solicita citas de entrega."),
            ("RF", "Requisito funcional."),
            ("RNF", "Requisito no funcional."),
        ]
        for term, meaning in terms:
            self.bullet(f"{term}: {meaning}")

        self.h("1.4 Referencias", 2)
        self.bullet("IEEE Std 830-1998 — IEEE Recommended Practice for Software Requirements Specifications.")

        self.h("1.5 Visión general del documento", 2)
        self._p(
            "La sección 2 describe el producto en términos generales. La sección 3 detalla "
            "requisitos específicos numerados (RF-xxx, RNF-xxx). La sección 4 resume trazabilidad "
            "entre actores y requisitos."
        )

    def _section_2_general(self) -> None:
        self.h("2. Descripción general", 1)

        self.h("2.1 Perspectiva del producto", 2)
        self._p(
            "Ferragro Citas es un sistema independiente que centraliza el agendamiento de entregas. "
            "Se despliega típicamente con frontend estático (Vercel), API (Render) y PostgreSQL (Render). "
            "Los proveedores interactúan por registro público o alta administrada; el personal interno "
            "gestiona franjas, estados de citas y usuarios."
        )

        self.h("2.2 Funciones del producto", 2)
        functions = [
            "Autenticación unificada (correo y contraseña) con roles Admin, Logística y Proveedor.",
            "Registro de proveedores con validación de NIT, dígito de verificación y datos de contacto.",
            "Creación y consulta de citas con validación de anticipación mínima y conflictos horarios.",
            "Configuración de franjas horarias por día y calendario operativo (zona America/Bogota).",
            "Revisión, extensión, reprogramación y cancelación de citas según rol.",
            "Gestión de usuarios internos y proveedores por personal autorizado.",
            "Notificaciones en aplicación, recordatorios programados y correo opcional (SMTP).",
            "Auditoría de inicios de sesión, eventos administrativos e historial de cambios en citas.",
            "Exportación de citas a hoja de cálculo y panel de analítica para personal interno.",
            "Perfil de usuario: actualización de datos, contraseña y foto (almacenamiento opcional en nube).",
        ]
        for f in functions:
            self.bullet(f)

        self.h("2.3 Características de los usuarios", 2)
        self._p("Clases de usuario y competencias esperadas:")
        self.bullet("Administrador: conocimiento del proceso logístico y del panel; configura franjas, usuarios y políticas operativas.")
        self.bullet("Logística: opera citas diarias, proveedores y reportes; no elimina proveedores ni citas salvo permisos explícitos de Admin donde aplique.")
        self.bullet("Proveedor: usuario de negocio con navegador; agenda y consulta sus propias citas; puede cancelar según reglas del sistema.")
        self.bullet("Visitante: accede a pantalla de inicio y registro/login sin autenticación previa.")

        self.h("2.4 Restricciones generales", 2)
        constraints = [
            "NIT de proveedor: exactamente 10 dígitos numéricos; dígito de verificación: 1 dígito.",
            "Documento de persona responsable del proveedor: entre 7 y 10 dígitos.",
            "Documento de usuario interno: solo dígitos; longitud según reglas de validación del API.",
            "Anticipación mínima para agendar o reprogramar (proveedor): configurable (por defecto 24 horas antes del inicio).",
            "Duración por defecto de una cita: 90 minutos, ampliable por personal interno.",
            "Estados de cita: sin revisión, revisado, finalizada, no presentada, cancelado.",
            "Un correo electrónico no puede asociarse a más de una credencial activa.",
            "Un NIT no puede registrarse dos veces mientras exista el proveedor en base de datos.",
            "Producción debe operar bajo HTTPS; cookies de actualización seguras en entorno productivo.",
        ]
        for c in constraints:
            self.bullet(c)

        self.h("2.5 Supuestos y dependencias", 2)
        self.bullet("Disponibilidad de PostgreSQL y conectividad de red entre cliente, API y base de datos.")
        self.bullet("Reloj del servidor sincronizado (UTC en base de datos; presentación en zona America/Bogota).")
        self.bullet("Para recuperación de contraseña por correo: servidor SMTP configurado en el API.")
        self.bullet("Para fotos de perfil en nube: cuenta Cloudinary configurada (opcional).")
        self.bullet("Los proveedores disponen de correo electrónico válido para notificaciones.")

    def _section_3_specific(self) -> None:
        self.h("3. Requisitos específicos", 1)
        self._p(
            "Los requisitos siguientes describen comportamiento observable del sistema. "
            "La prioridad Alta indica funcionalidad crítica para operación; Media, importante pero no bloqueante del día a día."
        )

        self.h("3.1 Requisitos funcionales", 2)

        self.h("3.1.1 Autenticación y sesión", 2)
        self.req_functional(
            "Registro de proveedor",
            "El visitante puede registrarse como proveedor indicando NIT, dígito de verificación, "
            "datos de empresa, persona responsable, correo y contraseña. El sistema rechaza NIT o correo duplicados.",
            actors="Visitante",
        )
        self.req_functional(
            "Inicio de sesión",
            "El usuario autentica con correo y contraseña. El sistema entrega token de acceso (JWT) "
            "y establece cookie HttpOnly de actualización. Registra intentos fallidos.",
            actors="Admin, Logística, Proveedor",
        )
        self.req_functional(
            "Bloqueo por intentos fallidos",
            "Tras superar el máximo de intentos fallidos (por defecto 5), la cuenta queda bloqueada "
            "durante un período configurable (por defecto 15 minutos).",
            priority="Alta",
        )
        self.req_functional(
            "Renovación de sesión",
            "Con cookie de actualización válida, el cliente obtiene un nuevo token de acceso sin volver a ingresar contraseña.",
            priority="Alta",
        )
        self.req_functional(
            "Cierre de sesión",
            "El usuario puede cerrar sesión; el sistema revoca la sesión de actualización y el cliente elimina el token local.",
        )
        self.req_functional(
            "Cierre en todos los dispositivos",
            "El usuario autenticado puede invalidar todas las sesiones de actualización asociadas a su credencial.",
            priority="Media",
        )
        self.req_functional(
            "Recuperación de contraseña",
            "Mediante correo registrado, el usuario solicita restablecimiento; el sistema aplica cooldown "
            "entre solicitudes y envía contraseña temporal por SMTP si está configurado.",
            priority="Alta",
        )
        self.req_functional(
            "Cambio de contraseña",
            "El usuario autenticado puede cambiar su contraseña proporcionando la actual y la nueva.",
        )

        self.h("3.1.2 Usuarios internos y roles", 2)
        self.req_functional(
            "Listado de roles",
            "Admin y Logística consultan los roles definidos en el sistema.",
            actors="Admin, Logística",
            priority="Media",
        )
        self.req_functional(
            "Administración de roles",
            "Admin crea, actualiza y elimina roles cuando no existan dependencias que lo impidan.",
            actors="Admin",
            priority="Media",
        )
        self.req_functional(
            "Alta de usuario interno",
            "Admin crea usuarios con documento, nombre, correo, contraseña y rol Admin o Logística. "
            "Libera correos huérfanos si aplica antes de crear.",
            actors="Admin",
        )
        self.req_functional(
            "Consulta y edición de usuarios internos",
            "Admin y Logística consultan usuarios; Admin actualiza datos y credenciales.",
            actors="Admin, Logística",
        )
        self.req_functional(
            "Eliminación de usuario interno",
            "Admin elimina un usuario interno y su credencial asociada, conservando historial de auditoría referenciado por documento.",
            actors="Admin",
        )
        self.req_functional(
            "Liberación de correo",
            "Admin ejecuta liberación de un correo eliminando credenciales huérfanas sin usuario ni proveedor activo.",
            actors="Admin",
            priority="Media",
        )

        self.h("3.1.3 Proveedores", 2)
        self.req_functional(
            "Listado y detalle de proveedores",
            "Admin y Logística consultan el catálogo de proveedores y el detalle por NIT.",
            actors="Admin, Logística",
        )
        self.req_functional(
            "Creación de proveedor por staff",
            "Admin o Logística registran un proveedor con credenciales; el sistema valida unicidad de NIT y correo.",
            actors="Admin, Logística",
        )
        self.req_functional(
            "Actualización de proveedor",
            "Admin o Logística modifican datos del proveedor y, si aplica, correo o contraseña.",
            actors="Admin, Logística",
        )
        self.req_functional(
            "Eliminación de proveedor",
            "Admin elimina un proveedor y su credencial si no existen dependencias (p. ej. citas) que lo impidan.",
            actors="Admin",
        )

        self.h("3.1.4 Citas de entrega", 2)
        self.req_functional(
            "Solicitud de cita por proveedor",
            "El proveedor crea una cita indicando material, fecha/hora de inicio y duración. "
            "El sistema valida anticipación mínima, franja permitida y ausencia de conflicto horario.",
            actors="Proveedor",
        )
        self.req_functional(
            "Consulta de citas por proveedor",
            "El proveedor consulta sus citas en modos lista, día o mes con filtros de estado.",
            actors="Proveedor",
        )
        self.req_functional(
            "Consulta operativa de citas",
            "Admin y Logística consultan citas con filtros por fecha, estado, proveedor y ordenamiento.",
            actors="Admin, Logística",
        )
        self.req_functional(
            "Creación de cita por Admin",
            "Admin puede crear citas en nombre de un proveedor seleccionado.",
            actors="Admin",
        )
        self.req_functional(
            "Actualización de cita",
            "Admin y Logística modifican datos de una cita existente según reglas de negocio.",
            actors="Admin, Logística",
        )
        self.req_functional(
            "Eliminación de cita",
            "Admin elimina una cita cuando la operación lo requiere.",
            actors="Admin",
            priority="Media",
        )
        self.req_functional(
            "Cambio de estado",
            "Personal autorizado cambia el estado (sin revisión, revisado, finalizada, no presentada, cancelado) con registro en historial.",
            actors="Admin, Logística, Proveedor (según operación)",
        )
        self.req_functional(
            "Extensión de duración",
            "Admin o Logística amplían la duración de una cita activa respetando disponibilidad.",
            actors="Admin, Logística",
        )
        self.req_functional(
            "Reprogramación",
            "Se reprograma una cita a nuevo horario validando franjas y conflictos.",
            actors="Admin, Logística, Proveedor (según flujo)",
        )
        self.req_functional(
            "Cancelación por proveedor",
            "El proveedor cancela su cita dentro de las reglas definidas (estados permitidos y plazos).",
            actors="Proveedor",
        )
        self.req_functional(
            "Verificación de conflicto",
            "Antes de confirmar horario, el sistema indica si existe solapamiento con otra cita del mismo proveedor.",
            priority="Alta",
        )
        self.req_functional(
            "Consulta de cupos disponibles",
            "El sistema calcula franjas disponibles para una fecha y proveedor según configuración y citas existentes.",
            actors="Proveedor, Admin, Logística",
        )
        self.req_functional(
            "Exportación de citas",
            "Admin y Logística descargan reporte de citas en formato de hoja de cálculo (.xlsx).",
            actors="Admin, Logística",
            priority="Media",
        )
        self.req_functional(
            "Finalización automática",
            "El sistema puede marcar citas como finalizadas cuando su ventana horaria ya concluyó, según tarea programada.",
            priority="Media",
        )

        self.h("3.1.5 Franjas horarias y calendario", 2)
        self.req_functional(
            "Consulta de franjas",
            "Usuarios autenticados consultan la configuración de franjas por día de la semana.",
            actors="Admin, Logística, Proveedor",
        )
        self.req_functional(
            "Franjas resueltas por fecha",
            "El sistema resuelve franjas efectivas para una fecha concreta (incluye excepciones por fecha).",
        )
        self.req_functional(
            "Configuración de franjas semanales",
            "Admin define o actualiza intervalos permitidos por día de la semana.",
            actors="Admin",
        )
        self.req_functional(
            "Excepciones por fecha",
            "Admin configura franjas o cierres específicos para fechas puntuales (festivos, mantenimiento).",
            actors="Admin",
            priority="Media",
        )
        self.req_functional(
            "Resumen de disponibilidad por fecha",
            "Se expone resumen de capacidad/ocupación por fecha para apoyo a la programación.",
            actors="Admin, Logística, Proveedor",
            priority="Media",
        )

        self.h("3.1.6 Perfil, notificaciones y auditoría", 2)
        self.req_functional(
            "Perfil propio",
            "Todo usuario autenticado consulta y actualiza su perfil (nombre, correo donde aplique).",
        )
        self.req_functional(
            "Foto de perfil",
            "El usuario sube o elimina foto de perfil; almacenamiento en servicio de medios en nube si está configurado.",
            priority="Media",
        )
        self.req_functional(
            "Centro de notificaciones",
            "El usuario recibe notificaciones en aplicación, las marca como leídas individual o masivamente.",
            priority="Media",
        )
        self.req_functional(
            "Recordatorios de cita",
            "El sistema envía recordatorios automáticos antes del inicio de citas en estados elegibles.",
            priority="Media",
        )
        self.req_functional(
            "Historial de cambios",
            "Admin y Logística consultan el historial de modificaciones de citas y acciones relevantes.",
            actors="Admin, Logística",
        )
        self.req_functional(
            "Auditoría de administración",
            "Admin consulta eventos administrativos (altas, bajas, cambios de configuración).",
            actors="Admin",
            priority="Media",
        )
        self.req_functional(
            "Auditoría de accesos",
            "El sistema registra intentos de login exitosos y fallidos para análisis de seguridad.",
            priority="Media",
        )

        self.h("3.1.7 Interfaz web (frontend)", 2)
        self.req_functional(
            "Página de inicio y acceso",
            "El visitante accede a landing, enlace de login y registro de proveedor.",
            actors="Visitante",
        )
        self.req_functional(
            "Panel por rol",
            "Tras autenticación, el usuario accede a un panel con menú acorde a su rol (citas, franjas, equipo, auditoría, analítica, configuración).",
        )
        self.req_functional(
            "Tema claro y oscuro",
            "El usuario alterna tema visual; la preferencia se conserva en el cliente.",
            priority="Media",
        )
        self.req_functional(
            "Tour guiado",
            "El panel ofrece recorrido guiado por módulos principales para facilitar adopción.",
            priority="Baja",
        )
        self.req_functional(
            "Mensajes de error de API",
            "La interfaz muestra mensajes devueltos por el API de forma comprensible al usuario final.",
            priority="Alta",
        )
        self.req_functional(
            "Renovación silenciosa de token",
            "Ante expiración del token de acceso, el cliente solicita renovación automática usando la cookie de actualización.",
            priority="Alta",
        )

        self.h("3.2 Requisitos no funcionales", 2)

        self.req_nonfunctional(
            "Rendimiento del API",
            "Las operaciones frecuentes de consulta y mutación deben completarse en menos de 500 ms "
            "en condiciones normales de producción (excluyendo arranque en frío del hosting gratuito).",
            category="Eficiencia",
            metric="p95 < 500 ms en /health y consultas de listado con paginación razonable.",
        )
        self.req_nonfunctional(
            "Tiempo de carga de interfaz",
            "Las pantallas principales del panel deben ser interactivas en menos de 2 segundos en red corporativa típica.",
            category="Eficiencia",
            metric="First meaningful paint < 2 s en conexión ≥ 10 Mbps.",
        )
        self.req_nonfunctional(
            "Seguridad de credenciales",
            "Las contraseñas se almacenan con hash bcrypt; nunca en texto plano.",
            category="Seguridad",
        )
        self.req_nonfunctional(
            "Control de acceso",
            "Cada endpoint exige autenticación y rol autorizado; respuestas 401/403 ante acceso indebido.",
            category="Seguridad",
        )
        self.req_nonfunctional(
            "Limitación de tasa",
            "El API aplica límites de peticiones por minuto (mayor restricción en rutas de autenticación).",
            category="Seguridad",
            metric="Por defecto 120 req/min general y 20 req/min en auth.",
        )
        self.req_nonfunctional(
            "Disponibilidad",
            "El servicio expone endpoint de salud; objetivo de disponibilidad mensual del 99 % en producción.",
            category="Confiabilidad",
            metric="GET /health responde 200 con identificador de build.",
        )
        self.req_nonfunctional(
            "Usabilidad",
            "Interfaz responsiva usable en escritorio y móvil; textos de validación en español.",
            category="Usabilidad",
        )
        self.req_nonfunctional(
            "Mantenibilidad",
            "Código modular (capas API, servicios, modelos); contrato OpenAPI publicado en /docs.",
            category="Mantenibilidad",
        )
        self.req_nonfunctional(
            "Trazabilidad de peticiones",
            "El API acepta o genera identificador de correlación (X-Correlation-ID) para soporte.",
            category="Operación",
        )
        self.req_nonfunctional(
            "Respuesta API uniforme",
            "Todas las respuestas JSON siguen el esquema: success, data, message.",
            category="Interoperabilidad",
        )
        self.req_nonfunctional(
            "Persistencia e integridad",
            "PostgreSQL garantiza unicidad de NIT, correo de proveedor y credencial; integridad referencial en citas.",
            category="Fiabilidad de datos",
        )
        self.req_nonfunctional(
            "Zona horaria operativa",
            "Validaciones de franja y calendario usan zona America/Bogota de forma consistente entre API e interfaz.",
            category="Exactitud",
        )

        self.h("3.3 Requisitos de interfaz externa", 2)

        self.h("3.3.1 Interfaces de usuario", 3)
        self.bullet("Pantallas: inicio, login, registro proveedor, panel principal, formularios de cita, calendario de franjas, gestión de equipo y proveedores, auditoría, analítica, configuración y perfil.")
        self.bullet("Navegadores objetivo: últimas versiones de Chrome, Edge y Firefox.")
        self.bullet("Resolución mínima recomendada: 360 px de ancho (móvil).")

        self.h("3.3.2 Interfaces de hardware", 3)
        self.bullet("Cliente: PC, tableta o smartphone con navegador y conexión a Internet.")
        self.bullet("Servidor: instancia compatible con ASGI (uvicorn) y PostgreSQL 14+.")

        self.h("3.3.3 Interfaces de software", 3)
        self.bullet("API REST JSON bajo prefijo /api/v1 (auth, crud, appointments, notifications, admin).")
        self.bullet("PostgreSQL como único almacén transaccional.")
        self.bullet("SMTP opcional para correos transaccionales (bienvenida, recuperación, recordatorios).")
        self.bullet("Cloudinary opcional para imágenes de perfil.")

        self.h("3.3.4 Interfaces de comunicaciones", 3)
        self.bullet("HTTPS obligatorio en producción entre cliente y API.")
        self.bullet("CORS configurado para orígenes del frontend desplegado.")
        self.bullet("Cookies de actualización: HttpOnly, Secure y SameSite acorde al entorno productivo.")

        self.h("3.4 Otros requisitos", 2)
        self.bullet("Idioma de la interfaz y mensajes al usuario: español (Colombia).")
        self.bullet("Protección de datos personales: minimizar exposición de contraseñas y tokens en logs.")
        self.bullet("Los usuarios y proveedores no se eliminan automáticamente por tiempo; la baja es acción explícita de Admin.")
        self.bullet("Reutilización de NIT o correo tras eliminación completa del registro previo y ausencia de bloqueos por citas dependientes.")

    def _section_4_traceability(self) -> None:
        self.doc.add_page_break()
        self.h("4. Matriz de trazabilidad resumida", 1)
        self._p("Relación entre actores principales y grupos de requisitos funcionales:")
        table = self.doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "Actor"
        hdr[1].text = "Grupos RF"
        hdr[2].text = "Requisitos clave"
        rows = [
            ("Visitante", "3.1.1, 3.1.7", "RF-001, registro e inicio"),
            ("Proveedor", "3.1.1, 3.1.4, 3.1.5, 3.1.6, 3.1.7", "Citas propias, cupos, perfil"),
            ("Logística", "3.1.2–3.1.6, 3.1.7", "Operación de citas y proveedores"),
            ("Admin", "Todos los grupos 3.1.x", "Configuración, eliminaciones, franjas, auditoría"),
        ]
        for actor, groups, key in rows:
            row = table.add_row().cells
            row[0].text = actor
            row[1].text = groups
            row[2].text = key

        self.doc.add_paragraph()
        self.h("Historial de revisiones", 2)
        rev = self.doc.add_table(rows=2, cols=4)
        rev.style = "Table Grid"
        rev.rows[0].cells[0].text = "Versión"
        rev.rows[0].cells[1].text = "Fecha"
        rev.rows[0].cells[2].text = "Autor"
        rev.rows[0].cells[3].text = "Descripción"
        rev.rows[1].cells[0].text = VERSION
        rev.rows[1].cells[1].text = TODAY
        rev.rows[1].cells[2].text = "Equipo Ferragro"
        rev.rows[1].cells[3].text = "Emisión inicial de la especificación IEEE 830."


def build_document() -> Path:
    b = DocBuilder()
    b.build().save(str(OUTPUT))
    return OUTPUT


if __name__ == "__main__":
    b = DocBuilder()
    out = b.build()
    out.save(str(OUTPUT))
    print(f"Documento generado: {OUTPUT}")
    print(f"Requisitos funcionales: {b._rf}, no funcionales: {b._rnf}")

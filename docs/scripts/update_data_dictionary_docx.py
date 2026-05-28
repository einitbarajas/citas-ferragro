# -*- coding: utf-8 -*-
"""Regenera docs/DICCIONARIO_DATOS_FERRAGRO.docx desde plantilla, una tabla por entidad."""

from __future__ import annotations

import shutil
from datetime import datetime
from pathlib import Path

from docx import Document

TEMPLATE_PATH = Path("C:/Users/ebarajas/Downloads/DICCIONARIO_DATOS_FERRAGRO.docx")

# (campo, tipo_sql, restricciones, descripcion)
Row = tuple[str, str, str, str]


def F(name: str) -> str:
    return f"`{name}`"


def T(typ: str) -> str:
    return f"`{typ}`"


SUMMARY_ROWS = [
    ("Credenciales", "Autenticacion unificada por correo y hash de contrasena", "IdCredencial"),
    ("Rol", "Catalogo de roles del portal", "Id"),
    ("Usuarios", "Personal interno con rol y credencial asociada", "IdDocumento"),
    ("UsuariosBodegas", "Bodegas asignadas al rol AdminBodega", "IdDocumento, IdBodega"),
    ("Proveedores", "Empresas proveedoras y contacto responsable", "IdNit"),
    ("Bodegas", "Lugares de entrega donde se agendan citas", "Id"),
    ("EquiposDescargaBodega", "Muelles o equipos de descarga por bodega", "Id"),
    ("Citas", "Citas de entrega de materiales", "Id"),
    ("HistorialCambios", "Auditoria de cambios sobre citas", "Id"),
    ("FranjasPermitidasCita", "Franjas horarias base por bodega y equipo", "Id"),
    ("FranjasPermitidasCitaFecha", "Franjas horarias especiales por fecha", "Id"),
    ("PerfilFoto", "Foto de perfil asociada a una credencial", "Id"),
    ("AuditoriaSistema", "Eventos administrativos del sistema", "Id"),
    ("SesionesRefresh", "Sesiones de refresh token por credencial", "Id"),
    ("IntentosLogin", "Control de intentos fallidos y bloqueo temporal", "IdCredencial"),
    ("AuditoriaLogin", "Auditoria de inicios de sesion exitosos o fallidos", "Id"),
    ("EjecucionesRecordatorio", "Trazabilidad de recordatorios automaticos", "Id"),
    ("EstadoResetContrasena", "Estado de restablecimiento y cambio obligatorio", "IdCredencial"),
    ("Notificaciones", "Notificaciones in-app por rol y proveedor", "Id"),
]

FK_ROWS = [
    ("Usuarios", "IdCredencial", "Credenciales", "N:1"),
    ("Usuarios", "IdRol", "Rol", "N:1"),
    ("UsuariosBodegas", "IdDocumento", "Usuarios", "N:1"),
    ("UsuariosBodegas", "IdBodega", "Bodegas", "N:1"),
    ("Proveedores", "IdCredencial", "Credenciales", "1:1"),
    ("EquiposDescargaBodega", "IdBodega", "Bodegas", "N:1"),
    ("Citas", "IdProveedor", "Proveedores", "N:1"),
    ("Citas", "IdBodega", "Bodegas", "N:1"),
    ("Citas", "IdEquipoDescargaBodega", "EquiposDescargaBodega", "N:1"),
    ("FranjasPermitidasCita", "IdBodega", "Bodegas", "N:1"),
    ("FranjasPermitidasCita", "IdEquipoDescargaBodega", "EquiposDescargaBodega", "N:1"),
    ("FranjasPermitidasCitaFecha", "IdBodega", "Bodegas", "N:1"),
    ("FranjasPermitidasCitaFecha", "IdEquipoDescargaBodega", "EquiposDescargaBodega", "N:1"),
    ("HistorialCambios", "IdCita", "Citas", "N:1"),
    ("PerfilFoto", "IdCredencial", "Credenciales", "1:1"),
    ("SesionesRefresh", "IdCredencial", "Credenciales", "N:1"),
    ("IntentosLogin", "IdCredencial", "Credenciales", "1:1"),
    ("AuditoriaLogin", "IdCredencial", "Credenciales", "N:1"),
    ("EstadoResetContrasena", "IdCredencial", "Credenciales", "1:1"),
    ("EjecucionesRecordatorio", "IdCita", "Citas", "N:1"),
    ("Notificaciones", "IdCita", "Citas", "N:1"),
    ("Notificaciones", "IdProveedorDestinatario", "Proveedores", "N:1"),
]

CONSTRAINT_ROWS = [
    ("Citas", "ENUM Estado", "Estados validos: sin_revision, revisado, finalizada, no_presentada, cancelado."),
    ("Citas", "FK IdProveedor / IdBodega / IdEquipoDescargaBodega", "Toda cita apunta a proveedor, bodega y equipo validos."),
    ("Citas", "DuracionMinutos DEFAULT 90", "Duracion operativa por defecto de 90 minutos."),
    ("Usuarios", "CHECK IdDocumento", "Documento interno con 7 a 10 digitos numericos."),
    ("Proveedores", "CHECK DigitoVerificacion", "Digito de verificacion de un solo digito."),
    ("Proveedores", "CHECK DocumentoPersonaResponsable", "Documento del responsable con 7 a 10 digitos."),
    ("Proveedores", "CHECK Estado", "Estado en {activo, suspendido}."),
    ("Bodegas", "CHECK EquiposDescarga", "Cantidad de equipos entre 1 y 20."),
    ("EquiposDescargaBodega", "UNIQUE (IdBodega, Nombre)", "No se repite el nombre del equipo en la misma bodega."),
    ("FranjasPermitidasCita", "CHECK HoraFin > HoraInicio", "La franja base debe tener hora fin posterior a inicio."),
    ("FranjasPermitidasCitaFecha", "UNIQUE por fecha/bodega/orden", "No se repite el orden de franja en un mismo dia."),
    ("Credenciales", "UNIQUE Correo", "Un correo solo puede autenticar una cuenta del sistema."),
    ("UsuariosBodegas", "PK compuesta", "Define alcance operativo del rol AdminBodega."),
    ("HistorialCambios", "Sin FK en IdActor", "Permite registrar actores internos o NIT de proveedor."),
]

DETAIL_BY_TABLE: dict[str, list[Row]] = {
    "Credenciales": [
        ("IdCredencial", "SERIAL", "PK, NOT NULL", "Identificador unico de credencial."),
        ("Correo", "VARCHAR(255)", "UNIQUE, NOT NULL", "Correo de acceso al portal."),
        ("HashContrasena", "VARCHAR(255)", "NOT NULL", "Contrasena almacenada con hash seguro."),
    ],
    "Rol": [
        ("Id", "SERIAL", "PK, NOT NULL", "Identificador del rol."),
        ("Nombre", "VARCHAR(40)", "UNIQUE, NOT NULL", "Nombre del rol (Admin, Logistica, AdminBodega, Proveedor)."),
    ],
    "Usuarios": [
        ("IdDocumento", "VARCHAR(30)", "PK, NOT NULL, CHECK", "Documento del usuario interno (7-10 digitos)."),
        ("NombreCompleto", "VARCHAR(120)", "NOT NULL", "Nombre completo del usuario."),
        ("IdCredencial", "INTEGER", "FK, UNIQUE, NOT NULL", "Credencial de acceso asociada."),
        ("IdRol", "INTEGER", "FK, NOT NULL", "Rol asignado en el portal."),
    ],
    "UsuariosBodegas": [
        ("IdDocumento", "VARCHAR(30)", "PK, FK, NOT NULL", "Usuario interno con alcance por bodega."),
        ("IdBodega", "INTEGER", "PK, FK, NOT NULL", "Bodega permitida para operar citas."),
    ],
    "Proveedores": [
        ("IdNit", "NUMERIC(10,0)", "PK, NOT NULL", "NIT de la empresa proveedora."),
        ("DigitoVerificacion", "VARCHAR(1)", "NOT NULL, CHECK", "Digito de verificacion del NIT."),
        ("NombreEmpresa", "VARCHAR(160)", "NOT NULL", "Razon social o nombre de la empresa."),
        ("CorreoEmpresa", "VARCHAR(255)", "UNIQUE, NOT NULL", "Correo corporativo de contacto."),
        ("IdCredencial", "INTEGER", "FK, UNIQUE, NOT NULL", "Credencial de acceso del proveedor."),
        ("NombrePersonaResponsable", "VARCHAR(160)", "NOT NULL", "Nombre del contacto responsable."),
        ("DocumentoPersonaResponsable", "VARCHAR(30)", "NOT NULL, CHECK", "Documento del contacto (7-10 digitos)."),
        ("Estado", "VARCHAR(20)", "NOT NULL, CHECK", "Estado de cuenta: activo o suspendido."),
        ("SuspendidoEn", "TIMESTAMPTZ", "NULL", "Fecha de suspension de la cuenta."),
        ("MotivoSuspension", "TEXT", "NULL", "Motivo registrado de la suspension."),
        ("SuspendidoPor", "VARCHAR(30)", "NULL", "Actor que suspendio la cuenta."),
        ("PurgaProgramadaEn", "TIMESTAMPTZ", "NULL", "Fecha programada de purga de datos."),
        ("EquiposDescarga", "INTEGER", "NOT NULL, CHECK 1..20", "Cantidad de equipos de descarga del proveedor."),
    ],
    "Bodegas": [
        ("Id", "SERIAL", "PK, NOT NULL", "Identificador de la bodega."),
        ("Nombre", "VARCHAR(120)", "UNIQUE, NOT NULL", "Nombre visible de la bodega."),
        ("Direccion", "VARCHAR(255)", "NULL", "Direccion fisica de la bodega."),
        ("Activa", "BOOLEAN", "NOT NULL, DEFAULT TRUE", "Indica si la bodega esta operativa."),
        ("Orden", "INTEGER", "NOT NULL, DEFAULT 0", "Orden de presentacion en listados."),
        ("EquiposDescarga", "INTEGER", "NOT NULL, CHECK 1..20", "Cantidad de equipos o muelles configurados."),
    ],
    "EquiposDescargaBodega": [
        ("Id", "SERIAL", "PK, NOT NULL", "Identificador del equipo de descarga."),
        ("IdBodega", "INTEGER", "FK, NOT NULL", "Bodega a la que pertenece el equipo."),
        ("Nombre", "VARCHAR(80)", "NOT NULL", "Nombre del muelle o equipo."),
        ("Activo", "BOOLEAN", "NOT NULL, DEFAULT TRUE", "Indica si el equipo esta disponible."),
        ("Orden", "INTEGER", "NOT NULL, DEFAULT 0", "Orden de presentacion dentro de la bodega."),
    ],
    "Citas": [
        ("Id", "SERIAL", "PK, NOT NULL", "Identificador de la cita."),
        ("IdProveedor", "NUMERIC(10,0)", "FK, NOT NULL", "Proveedor que agenda la entrega."),
        ("IdBodega", "INTEGER", "FK, NOT NULL", "Bodega donde se realiza la entrega."),
        ("IdEquipoDescargaBodega", "INTEGER", "FK, NOT NULL", "Muelle o equipo asignado."),
        ("IndiceEquipoProveedor", "INTEGER", "NOT NULL, DEFAULT 1", "Indice del equipo del proveedor en la cita."),
        ("DescripcionMaterial", "TEXT", "NOT NULL", "Descripcion del material a entregar."),
        ("FechaHoraInicio", "TIMESTAMPTZ", "NOT NULL", "Fecha y hora de inicio de la cita."),
        ("DuracionMinutos", "INTEGER", "NOT NULL, DEFAULT 90", "Duracion estimada en minutos."),
        ("Estado", "EstadoCita", "NOT NULL, DEFAULT sin_revision", "Estado operativo de la cita."),
    ],
    "HistorialCambios": [
        ("Id", "SERIAL", "PK, NOT NULL", "Identificador del evento de historial."),
        ("IdActor", "VARCHAR(30)", "NOT NULL", "Actor que realizo el cambio (documento o NIT)."),
        ("IdCita", "INTEGER", "FK, NOT NULL", "Cita afectada por el cambio."),
        ("Accion", "VARCHAR(80)", "NOT NULL", "Tipo de accion registrada."),
        ("Descripcion", "TEXT", "NOT NULL", "Detalle legible del cambio."),
        ("CreadoEn", "TIMESTAMPTZ", "NOT NULL", "Marca de tiempo del evento."),
        ("CampoCritico", "VARCHAR(80)", "NULL", "Campo sensible modificado."),
        ("ValorAnterior", "TEXT", "NULL", "Valor anterior del campo critico."),
        ("ValorNuevo", "TEXT", "NULL", "Valor nuevo del campo critico."),
    ],
    "FranjasPermitidasCita": [
        ("Id", "SERIAL", "PK, NOT NULL", "Identificador de la franja semanal."),
        ("IdBodega", "INTEGER", "FK, NOT NULL", "Bodega donde aplica la franja."),
        ("IdEquipoDescargaBodega", "INTEGER", "FK, NULL", "Equipo especifico (opcional)."),
        ("HoraInicio", "TIME", "NOT NULL", "Hora de inicio permitida."),
        ("HoraFin", "TIME", "NOT NULL, CHECK > inicio", "Hora de fin permitida."),
        ("Orden", "INTEGER", "NOT NULL, DEFAULT 0", "Orden del turno en la interfaz."),
    ],
    "FranjasPermitidasCitaFecha": [
        ("Id", "SERIAL", "PK, NOT NULL", "Identificador de la franja por fecha."),
        ("IdBodega", "INTEGER", "FK, NOT NULL", "Bodega donde aplica la franja."),
        ("IdEquipoDescargaBodega", "INTEGER", "FK, NULL", "Equipo especifico (opcional)."),
        ("Fecha", "DATE", "NOT NULL", "Dia calendario de la excepcion."),
        ("HoraInicio", "TIME", "NOT NULL", "Hora de inicio permitida."),
        ("HoraFin", "TIME", "NOT NULL, CHECK > inicio", "Hora de fin permitida."),
        ("Orden", "INTEGER", "NOT NULL, DEFAULT 0", "Orden del turno para ese dia."),
    ],
    "PerfilFoto": [
        ("Id", "SERIAL", "PK, NOT NULL", "Identificador del registro de foto."),
        ("IdCredencial", "INTEGER", "FK, UNIQUE, NOT NULL", "Credencial duena de la foto."),
        ("FotoUrl", "VARCHAR(500)", "NULL", "URL de la imagen de perfil."),
        ("ActualizadoEn", "TIMESTAMPTZ", "NOT NULL, DEFAULT NOW()", "Ultima actualizacion de la foto."),
    ],
    "AuditoriaSistema": [
        ("Id", "SERIAL", "PK, NOT NULL", "Identificador del evento administrativo."),
        ("IdActor", "VARCHAR(30)", "NOT NULL", "Documento del administrador que actua."),
        ("Accion", "VARCHAR(80)", "NOT NULL", "Accion ejecutada."),
        ("Descripcion", "TEXT", "NOT NULL", "Detalle del evento."),
        ("CreadoEn", "TIMESTAMPTZ", "NOT NULL", "Fecha del evento."),
        ("DocumentoObjetivo", "VARCHAR(30)", "NULL", "Documento o entidad impactada."),
    ],
    "SesionesRefresh": [
        ("Id", "SERIAL", "PK, NOT NULL", "Identificador de la sesion refresh."),
        ("IdCredencial", "INTEGER", "FK, NOT NULL", "Credencial asociada a la sesion."),
        ("Jti", "UUID", "UNIQUE, NOT NULL", "Identificador unico del token refresh."),
        ("CreadoEn", "TIMESTAMPTZ", "NOT NULL, DEFAULT NOW()", "Fecha de creacion de la sesion."),
        ("ExpiraEn", "TIMESTAMPTZ", "NOT NULL", "Fecha de expiracion del token."),
        ("RevocadoEn", "TIMESTAMPTZ", "NULL", "Fecha de revocacion (logout)."),
    ],
    "IntentosLogin": [
        ("IdCredencial", "INTEGER", "PK, FK, NOT NULL", "Credencial controlada por politica de bloqueo."),
        ("FallosConsecutivos", "INTEGER", "NOT NULL, DEFAULT 0", "Contador de intentos fallidos seguidos."),
        ("BloqueadoHasta", "TIMESTAMPTZ", "NULL", "Fin del bloqueo temporal por seguridad."),
    ],
    "AuditoriaLogin": [
        ("Id", "SERIAL", "PK, NOT NULL", "Identificador del intento auditado."),
        ("IdCredencial", "INTEGER", "FK, NULL", "Credencial asociada si existe."),
        ("Correo", "VARCHAR(255)", "NOT NULL", "Correo usado en el intento."),
        ("Exito", "BOOLEAN", "NOT NULL", "Indica si el login fue exitoso."),
        ("DireccionIp", "VARCHAR(45)", "NULL", "Direccion IP de origen."),
        ("UserAgent", "TEXT", "NULL", "Agente de usuario del navegador."),
        ("MotivoFallo", "VARCHAR(255)", "NULL", "Causa del fallo de autenticacion."),
        ("CreadoEn", "TIMESTAMPTZ", "NOT NULL, DEFAULT NOW()", "Fecha del intento."),
    ],
    "EjecucionesRecordatorio": [
        ("Id", "SERIAL", "PK, NOT NULL", "Identificador de la ejecucion."),
        ("IdCita", "INTEGER", "FK, NOT NULL", "Cita procesada por el recordatorio."),
        ("Tipo", "VARCHAR(40)", "NOT NULL, DEFAULT recordatorio_proximo", "Tipo de recordatorio enviado."),
        ("Estado", "VARCHAR(30)", "NOT NULL", "Resultado de la ejecucion."),
        ("Detalle", "TEXT", "NULL", "Detalle tecnico o mensaje de error."),
        ("EjecutadoEn", "TIMESTAMPTZ", "NOT NULL, DEFAULT NOW()", "Fecha de ejecucion del job."),
    ],
    "EstadoResetContrasena": [
        ("IdCredencial", "INTEGER", "PK, FK, NOT NULL", "Credencial con estado de reseteo."),
        ("DebeCambiarContrasena", "BOOLEAN", "NOT NULL", "Obliga cambio de clave al siguiente login."),
        ("EmitidoTemporalEn", "TIMESTAMPTZ", "NULL", "Fecha de emision de clave temporal."),
    ],
    "Notificaciones": [
        ("Id", "SERIAL", "PK, NOT NULL", "Identificador de la notificacion."),
        ("RolDestinatario", "VARCHAR(20)", "NOT NULL", "Rol destinatario de la notificacion."),
        ("IdProveedorDestinatario", "NUMERIC(10,0)", "NULL", "NIT del proveedor destinatario (si aplica)."),
        ("IdCita", "INTEGER", "FK, NOT NULL", "Cita relacionada con el aviso."),
        ("Tipo", "VARCHAR(40)", "NOT NULL", "Tipo de evento notificado."),
        ("Titulo", "VARCHAR(160)", "NOT NULL", "Titulo mostrado en el centro de notificaciones."),
        ("Mensaje", "TEXT", "NOT NULL", "Cuerpo del mensaje."),
        ("LeidaEn", "TIMESTAMPTZ", "NULL", "Fecha en que el usuario marco como leida."),
        ("CreadaEn", "TIMESTAMPTZ", "NOT NULL", "Fecha de creacion de la notificacion."),
    ],
}

ENUM_ESTADO_CITA: list[Row] = [
    ("sin_revision", "EstadoCita", "Valor por defecto al crear", "Cita creada, pendiente de revision."),
    ("revisado", "EstadoCita", "Valor del enum", "Cita confirmada por logistica."),
    ("finalizada", "EstadoCita", "Valor del enum", "Entrega completada."),
    ("no_presentada", "EstadoCita", "Valor del enum", "Proveedor no asistio en la ventana."),
    ("cancelado", "EstadoCita", "Valor del enum", "Cita cancelada por usuario autorizado."),
]

DETAIL_ORDER = [name for name, _, _ in SUMMARY_ROWS]


def clear_document_content(doc: Document) -> None:
    body = doc._element.body
    for child in list(body):
        if child.tag.endswith("sectPr"):
            continue
        body.remove(child)


def fmt_rows(rows: list[Row]) -> list[tuple[str, str, str, str]]:
    return [(F(c), T(t), r, d) for c, t, r, d in rows]


def add_grid_table(doc: Document, headers: list[str], rows: list[tuple[str, ...]], style: str) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = style
    for i, header in enumerate(headers):
        table.cell(0, i).text = header
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)


def add_entity_table(doc: Document, rows: list[Row], style: str) -> None:
    add_grid_table(
        doc,
        ["Campo", "Tipo", "Restricciones", "Descripcion"],
        fmt_rows(rows),
        style,
    )


def main() -> None:
    repo_root = Path(__file__).resolve().parents[2]
    out_path = repo_root / "docs" / "DICCIONARIO_DATOS_FERRAGRO.docx"

    if not TEMPLATE_PATH.exists():
        raise FileNotFoundError(f"No se encontro la plantilla: {TEMPLATE_PATH}")

    ref = Document(TEMPLATE_PATH)
    table_style = ref.tables[0].style.name if ref.tables[0].style else "Medium Shading 1 Accent 3"

    out_path.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(TEMPLATE_PATH, out_path)
    doc = Document(out_path)
    clear_document_content(doc)

    doc.add_paragraph("Diccionario de Datos ? Ferragro")
    doc.add_paragraph("")
    doc.add_paragraph("1) Tablas de la base de datos (resumen)")
    add_grid_table(doc, ["Tabla", "Proposito", "PK"], SUMMARY_ROWS, table_style)
    doc.add_paragraph("")
    doc.add_paragraph("2) Relaciones principales (FK)")
    add_grid_table(
        doc,
        ["Tabla hija", "Campo FK", "Tabla padre", "Cardinalidad"],
        FK_ROWS,
        table_style,
    )
    doc.add_paragraph("")
    doc.add_paragraph("3) Restricciones clave de negocio")
    add_grid_table(doc, ["Tabla", "Restriccion", "Regla"], CONSTRAINT_ROWS, table_style)
    doc.add_paragraph("")
    doc.add_paragraph(
        f"Documento actualizado: {datetime.now().strftime('%Y-%m-%d %H:%M')}. "
        "Basado en db/init/*.sql y modelos ORM del backend."
    )
    doc.add_paragraph("")

    for index, table_name in enumerate(DETAIL_ORDER, start=1):
        doc.add_paragraph(f"{index}) `{table_name}`")
        add_entity_table(doc, DETAIL_BY_TABLE[table_name], table_style)
        doc.add_paragraph("")

    doc.add_paragraph("Tipo enumerado `EstadoCita`")
    add_grid_table(
        doc,
        ["Valor", "Tipo", "Restricciones", "Descripcion"],
        fmt_rows(ENUM_ESTADO_CITA),
        table_style,
    )

    doc.save(out_path)
    print(f"Documento actualizado: {out_path}")
    print(f"Tablas detalladas: {len(DETAIL_ORDER)} + enum EstadoCita")


if __name__ == "__main__":
    main()

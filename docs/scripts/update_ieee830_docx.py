from datetime import datetime
from pathlib import Path

from docx import Document


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def main() -> None:
    repo_root = Path(__file__).resolve().parents[2]
    out_path = repo_root / "docs" / "ESPECIFICACION_REQUISITOS_IEEE830_FERRAGRO.docx"

    doc = Document()
    doc.add_heading("Especificacion de Requisitos de Software (IEEE 830)", 0)
    doc.add_paragraph("Proyecto: Ferragro - Gestion de Citas de Entrega")
    doc.add_paragraph(
        f"Fecha de actualizacion: {datetime.now().strftime('%Y-%m-%d %H:%M')}"
    )

    doc.add_heading("1. Introduccion", level=1)
    doc.add_heading("1.1 Proposito", level=2)
    doc.add_paragraph(
        "Este documento define los requisitos funcionales y no funcionales del portal "
        "Ferragro para agendamiento y operacion de citas de entrega, alineado con el "
        "estado actual del repositorio y despliegue en produccion."
    )
    doc.add_heading("1.2 Alcance", level=2)
    doc.add_paragraph(
        "El sistema permite registrar, consultar, reprogramar, revisar y auditar citas "
        "de entrega por bodega y muelle/equipo de descarga, con notificaciones in-app "
        "y por correo, autenticacion JWT con refresh token y tareas automaticas en segundo plano."
    )

    doc.add_heading("2. Descripcion general", level=1)
    doc.add_heading("2.1 Usuarios y roles", level=2)
    add_bullets(
        doc,
        [
            "Proveedor: crea y consulta sus citas.",
            "Logistica: opera citas y estados en tablero interno.",
            "AdminBodega: opera citas limitado a bodegas asignadas.",
            "Admin: administracion global, auditoria y configuraciones.",
        ],
    )
    doc.add_heading("2.2 Entorno de operacion", level=2)
    add_bullets(
        doc,
        [
            "Frontend (Vite/React): https://frontend-ferragro.vercel.app",
            "Backend API (FastAPI): https://ferragro-api.onrender.com",
            "Swagger: https://ferragro-api.onrender.com/docs",
            "Base de datos: PostgreSQL (local y Render).",
        ],
    )

    doc.add_heading("3. Requisitos funcionales", level=1)
    add_bullets(
        doc,
        [
            "RF-01: El sistema debe permitir autenticacion por correo y contrasena.",
            "RF-02: El sistema debe emitir access token y refresh token (cookie HttpOnly).",
            "RF-03: El proveedor debe poder crear citas con validacion de anticipacion minima.",
            "RF-04: El sistema debe validar conflictos de horario por bodega y muelle/equipo.",
            "RF-05: El personal interno debe listar citas por modos de consulta y filtros.",
            "RF-06: El personal autorizado debe actualizar estado de cita y registrar historial.",
            "RF-07: El personal autorizado debe poder extender una cita sin solapamiento.",
            "RF-08: El sistema debe enviar notificaciones in-app y, si aplica, correo.",
            "RF-09: El sistema debe exponer centro de notificaciones (listar, marcar leida, limpiar).",
            "RF-10: El sistema debe soportar rol AdminBodega con restriccion por UsuariosBodegas.",
            "RF-11: El sistema debe registrar trazabilidad en HistorialCambios.",
            "RF-12: El modulo admin debe permitir consulta de logs del sistema.",
        ],
    )

    doc.add_heading("4. Requisitos no funcionales", level=1)
    add_bullets(
        doc,
        [
            "RNF-01 Seguridad: uso de JWT, refresh token HttpOnly y politicas de CORS.",
            "RNF-02 Disponibilidad: operacion en Render/Vercel con verificacion de health.",
            "RNF-03 Mantenibilidad: CI automatizado en push/PR para backend y frontend.",
            "RNF-04 Rendimiento: validaciones de agenda por franja y equipo de descarga.",
            "RNF-05 Escalabilidad: separacion frontend/backend y BD administrada en la nube.",
            "RNF-06 Trazabilidad: auditoria de cambios y notificaciones persistidas.",
            "RNF-07 Compatibilidad: ejecucion local en Windows con PowerShell y despliegue cloud.",
        ],
    )

    doc.add_heading("5. Interfaces externas", level=1)
    doc.add_heading("5.1 API principal", level=2)
    add_bullets(
        doc,
        [
            "Auth: /api/auth/register, /api/auth/login, /api/auth/refresh, /api/auth/logout",
            "CRUD principal: /api/crud/... (citas, bodegas, franjas, usuarios, historial).",
            "Citas legacy: /api/appointments (listado, estado, extension).",
            "Notificaciones: /api/v1/notifications",
            "Administracion: /api/admin/logs",
        ],
    )
    doc.add_heading("5.2 Procesos en segundo plano", level=2)
    add_bullets(
        doc,
        [
            "reminder_scheduler: recordatorios de citas proximas.",
            "no_presentada_scheduler: marca no presentada con ventana de gracia.",
            "notification_purge_scheduler: elimina notificaciones vencidas.",
            "provider_purge_scheduler: purga proveedores suspendidos.",
        ],
    )

    doc.add_heading("6. Restricciones y reglas de negocio", level=1)
    add_bullets(
        doc,
        [
            "Zona horaria de negocio: BUSINESS_TIMEZONE (por defecto America/Bogota).",
            "Anticipacion minima para crear cita: APPOINTMENT_MINIMUM_NOTICE_HOURS.",
            "Cancelacion minima configurable: APPOINTMENT_CANCEL_MINIMUM_NOTICE_HOURS.",
            "Retencion de notificaciones: NOTIFICATION_RETENTION_DAYS.",
            "Operaciones staff sujetas a rol y alcance de bodega.",
        ],
    )

    doc.add_heading("7. Validacion y pruebas", level=1)
    doc.add_paragraph(
        "La calidad se valida con suite automatizada de backend (pytest), "
        "verificaciones operativas, build del frontend y pipeline CI en GitHub Actions."
    )

    doc.add_heading("8. Despliegue", level=1)
    add_bullets(
        doc,
        [
            "Frontend en Vercel con VITE_API_URL apuntando al backend.",
            "Backend en Render con health check /health.",
            "Base PostgreSQL en Render con migraciones db/init y funciones CRUD.",
            "Variables de entorno obligatorias para seguridad, DB y correo SMTP.",
        ],
    )

    doc.add_heading("9. Referencias", level=1)
    add_bullets(
        doc,
        [
            "README.md (fuente tecnica principal del repositorio).",
            "docs/GUIA_OPERACION_PRODUCCION.docx.",
            "docs/PRUEBAS.docx.",
            "db/README.md.",
            ".github/workflows/ci.yml.",
        ],
    )

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    print(f"Documento actualizado: {out_path}")


if __name__ == "__main__":
    main()

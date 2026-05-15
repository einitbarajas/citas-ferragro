"""Regenera GUIA_OPERACION_PRODUCCION.docx desde el .md (uso interno)."""
from pathlib import Path

from docx import Document
from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[1]
MD_PATH = ROOT / "GUIA_OPERACION_PRODUCCION.md"
DOCX_PATH = ROOT / "GUIA_OPERACION_PRODUCCION.docx"


def add_paragraph(doc: Document, text: str, style: str | None = None) -> None:
    p = doc.add_paragraph(text, style=style)
    for run in p.runs:
        run.font.size = Pt(11)


def main() -> None:
    text = MD_PATH.read_text(encoding="utf-8")
    doc = Document()
    doc.add_heading("Guía de operación — Ferragro Citas (producción)", 0)

    in_code = False
    code_lines: list[str] = []

    for raw in text.splitlines():
        line = raw.rstrip()

        if line.startswith("```"):
            if in_code:
                add_paragraph(doc, "\n".join(code_lines), style=None)
                code_lines = []
                in_code = False
            else:
                in_code = True
            continue

        if in_code:
            code_lines.append(line)
            continue

        if not line.strip():
            continue
        if line == "---":
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
            continue
        if line.startswith("|") and line.endswith("|"):
            add_paragraph(doc, line.replace("|", " ").strip())
            continue
        if line.startswith("> "):
            add_paragraph(doc, line[2:].strip())
            continue
        if line.startswith("- "):
            add_paragraph(doc, line[2:].strip(), style="List Bullet")
            continue
        add_paragraph(doc, line)

    doc.save(DOCX_PATH)
    print(f"Written {DOCX_PATH}")


if __name__ == "__main__":
    main()
